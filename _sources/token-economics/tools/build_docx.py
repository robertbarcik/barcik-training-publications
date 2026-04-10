#!/usr/bin/env python3
"""Build a DOCX booklet from chapter markdown sources using python-docx."""

import os
import re
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CHAPTERS_DIR = os.path.join(BASE_DIR, "chapters")
OUTPUT_FILE = os.path.join(BASE_DIR, "output", "booklet.docx")

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0x3B, 0x82, 0xF6)
DARK = RGBColor(0x1E, 0x29, 0x3B)
LIGHT = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_BG = RGBColor(0xF8, 0xFA, 0xFC)


def get_chapter_files():
    return sorted(glob.glob(os.path.join(CHAPTERS_DIR, "*.md")))


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear"
    })
    shading.append(shading_elem)


def add_formatted_text(paragraph, text, bold=False, italic=False, code=False, color=None):
    """Add a run with formatting to a paragraph."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color
    return run


def parse_inline(paragraph, text, default_bold=False):
    """Parse inline markdown (bold, italic, code) and add runs to paragraph."""
    # Process inline formatting: **bold**, *italic*, `code`
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))"
    for match in re.finditer(pattern, text):
        full = match.group(0)
        if match.group(2):  # **bold**
            add_formatted_text(paragraph, match.group(2), bold=True)
        elif match.group(3):  # *italic*
            add_formatted_text(paragraph, match.group(3), italic=True)
        elif match.group(4):  # `code`
            add_formatted_text(paragraph, match.group(4), code=True)
        elif match.group(5):  # plain text
            run = paragraph.add_run(match.group(5))
            if default_bold:
                run.bold = True


def add_table(doc, header_row, data_rows):
    """Add a formatted table to the document."""
    num_cols = len(header_row)
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for i, text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text.strip())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        set_cell_shading(cell, "1E3A5F")

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        row = table.add_row()
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            parse_inline(p, text.strip())
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.name = "Calibri"
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F8FAFC")

    doc.add_paragraph()  # spacing after table


def process_markdown(doc, content, is_frontmatter=False):
    """Process markdown content and add to docx document."""
    lines = content.split("\n")
    i = 0
    in_code_block = False
    in_table = False
    table_header = None
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            i += 1
            continue

        # Table detection
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Check if this is a separator row
            if all(re.match(r"^[-:]+$", c) for c in cells):
                i += 1
                continue
            if table_header is None:
                table_header = cells
                in_table = True
            else:
                table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            add_table(doc, table_header, table_rows)
            table_header = None
            table_rows = []
            in_table = False

        # Headings
        h1_match = re.match(r"^#\s+(.+)", line)
        h2_match = re.match(r"^##\s+(.+)", line)
        h3_match = re.match(r"^###\s+(.+)", line)

        if h1_match:
            heading = doc.add_heading(level=1)
            run = heading.add_run(h1_match.group(1).strip())
            run.font.color.rgb = NAVY
            run.font.size = Pt(26)
            i += 1
            continue

        if h2_match:
            heading = doc.add_heading(level=2)
            run = heading.add_run(h2_match.group(1).strip())
            run.font.color.rgb = NAVY
            run.font.size = Pt(18)
            i += 1
            continue

        if h3_match:
            heading = doc.add_heading(level=3)
            run = heading.add_run(h3_match.group(1).strip())
            run.font.color.rgb = NAVY
            run.font.size = Pt(13)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("* * *")
            run.font.color.rgb = LIGHT
            i += 1
            continue

        # Blockquote
        if line.strip().startswith(">"):
            text = line.strip().lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            parse_inline(p, text)
            for run in p.runs:
                run.font.color.rgb = NAVY
            i += 1
            continue

        # Unordered list
        ul_match = re.match(r"^(\s*)[-*]\s+(.+)", line)
        if ul_match:
            indent = len(ul_match.group(1))
            text = ul_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 2:
                p.paragraph_format.left_indent = Cm(1.5 + indent * 0.3)
            parse_inline(p, text)
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r"^(\s*)\d+\.\s+(.+)", line)
        if ol_match:
            text = ol_match.group(2)
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, text)
            i += 1
            continue

        # Empty line
        if line.strip() == "":
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline(p, line)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"

        i += 1

    # Flush any remaining table
    if in_table and table_header:
        add_table(doc, table_header, table_rows)


def build():
    """Build the DOCX file."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    files = get_chapter_files()

    for file_idx, f in enumerate(files):
        with open(f, "r", encoding="utf-8") as fh:
            content = fh.read().strip()

        # Add page break between chapters (not before the first one)
        if file_idx > 0:
            doc.add_page_break()

        is_frontmatter = "00_frontmatter" in f
        process_markdown(doc, content, is_frontmatter)

    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)

    print(f"Built: {OUTPUT_FILE}")
    print(f"  Chapters: {len(files)}")


if __name__ == "__main__":
    build()
